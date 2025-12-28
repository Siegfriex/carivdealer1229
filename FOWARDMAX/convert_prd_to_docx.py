#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
PRD_Phase1_2025-12-31.md를 DOCX로 변환하는 스크립트
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx 라이브러리가 필요합니다.")
    print("설치 방법: pip install python-docx")
    sys.exit(1)

def parse_markdown_table(line):
    """마크다운 테이블 행 파싱"""
    if '|' not in line:
        return None
    cells = [cell.strip() for cell in line.split('|')]
    # 첫 번째와 마지막이 비어있으면 제거
    if cells and not cells[0]:
        cells = cells[1:]
    if cells and not cells[-1]:
        cells = cells[:-1]
    return cells

def is_table_separator(line):
    """테이블 구분선인지 확인"""
    return re.match(r'^\s*\|[\s\-:]+\|\s*$', line)

def add_hyperlink(paragraph, text, url):
    """하이퍼링크 추가"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # 링크 스타일
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def convert_md_to_docx(md_file_path, docx_file_path):
    """마크다운 파일을 DOCX로 변환"""
    
    # DOCX 문서 생성
    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    
    # 제목 스타일 설정
    for heading_level in range(1, 7):
        heading_style = doc.styles[f'Heading {heading_level}']
        heading_font = heading_style.font
        heading_font.name = '맑은 고딕'
        heading_font.bold = True
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table = None
    table_header = None
    
    while i < len(lines):
        line = lines[i].rstrip('\n\r')
        
        # 빈 줄
        if not line.strip():
            if in_table:
                in_table = False
                table = None
                table_header = None
            else:
                doc.add_paragraph()
            i += 1
            continue
        
        # 제목 처리 (#, ##, ### 등)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            if in_table:
                in_table = False
                table = None
                table_header = None
            
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # 볼드 제거
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            
            if level == 1:
                para = doc.add_heading(text, level=1)
            elif level == 2:
                para = doc.add_heading(text, level=2)
            elif level == 3:
                para = doc.add_heading(text, level=3)
            elif level == 4:
                para = doc.add_heading(text, level=4)
            else:
                para = doc.add_heading(text, level=5)
            
            i += 1
            continue
        
        # 테이블 처리
        if '|' in line and not is_table_separator(line):
            cells = parse_markdown_table(line)
            if cells:
                if not in_table:
                    table = doc.add_table(rows=0, cols=len(cells))
                    table.style = 'Light Grid Accent 1'
                    in_table = True
                    table_header = cells
                    # 헤더 행 추가
                    header_row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        cell = header_row.cells[j]
                        cell.text = cell_text
                        # 헤더 셀 볼드 처리
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                else:
                    # 데이터 행 추가
                    row = table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
            i += 1
            continue
        
        # 테이블 구분선 무시
        if is_table_separator(line):
            i += 1
            continue
        
        # 인용구 처리 (>)
        if line.strip().startswith('>'):
            if in_table:
                in_table = False
                table = None
                table_header = None
            
            quote_text = line.replace('>', '').strip()
            para = doc.add_paragraph(quote_text)
            para.style = 'Intense Quote'
            i += 1
            continue
        
        # 코드 블록 처리 (```)
        if line.strip().startswith('```'):
            if in_table:
                in_table = False
                table = None
                table_header = None
            
            # 코드 블록 시작
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 닫는 ``` 건너뛰기
            
            if code_lines:
                code_text = ''.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'No Spacing'
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            continue
        
        # 인라인 코드 처리 (`)
        if '`' in line:
            if in_table:
                in_table = False
                table = None
                table_header = None
            
            para = doc.add_paragraph()
            parts = re.split(r'(`[^`]+`)', line)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    para.add_run(part)
            i += 1
            continue
        
        # 링크 처리 [텍스트](URL)
        if '[' in line and '](' in line:
            if in_table:
                in_table = False
                table = None
                table_header = None
            
            para = doc.add_paragraph()
            # 링크 패턴 찾기
            link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
            last_pos = 0
            for match in re.finditer(link_pattern, line):
                # 링크 앞 텍스트
                if match.start() > last_pos:
                    para.add_run(line[last_pos:match.start()])
                # 링크 추가
                link_text = match.group(1)
                link_url = match.group(2)
                add_hyperlink(para, link_text, link_url)
                last_pos = match.end()
            # 나머지 텍스트
            if last_pos < len(line):
                para.add_run(line[last_pos:])
            i += 1
            continue
        
        # 일반 텍스트
        if in_table:
            in_table = False
            table = None
            table_header = None
        
        # 볼드/이탤릭 처리
        text = line
        para = doc.add_paragraph()
        
        # 볼드 처리 **text**
        parts = re.split(r'(\*\*[^\*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.font.bold = True
            else:
                para.add_run(part)
        
        i += 1
    
    # 문서 저장
    doc.save(docx_file_path)
    print(f"변환 완료: {docx_file_path}")

if __name__ == '__main__':
    input_file = 'PRD_Phase1_2025-12-31.md'
    output_file = 'PRD_Phase1_2025-12-31_최종.docx'
    
    if not Path(input_file).exists():
        print(f"오류: {input_file} 파일을 찾을 수 없습니다.")
        sys.exit(1)
    
    print(f"마크다운 파일 읽는 중: {input_file}")
    convert_md_to_docx(input_file, output_file)
    print(f"DOCX 파일 생성 완료: {output_file}")

